# Copyright The Marin Authors
# SPDX-License-Identifier: Apache-2.0

import io
import json
import zipfile
from dataclasses import dataclass

import pytest
from docx import Document
from marin.datakit.download.common_crawl_docx import (
    DOCX_MIME_TYPE,
    CommonCrawlDocxConfig,
    DoclingDocxExtractor,
    DocumentBlock,
    DocxExtractionError,
    DocxRecordSelector,
    DocxSelectionReason,
    ExtractedDocument,
    InvalidDocxError,
    LinguaLanguageDetector,
    common_crawl_docx_steps,
    detect_document_language,
    extracted_docx_record,
    fetched_docx_record,
    language_chunks,
    language_identified_docx_record,
    process_fetched_docx,
    validate_docx,
)
from marin.datakit.download.common_crawl_plan import (
    CommonCrawlIndexKind,
    CommonCrawlSource,
    FetchedCommonCrawlRecord,
)
from marin.datakit.download.common_crawl_warc import CommonCrawlWarcRecord, content_digest, main_record_from_index_row

from experiments.datakit.common_crawl_docx_sample import sample_report_markdown

CRAWL_ID = "CC-MAIN-2026-30"
RECORD_ID = "<urn:uuid:019f8700-d21d-78d8-8eb1-99eaa22579da>"
URL = "https://example.com/report.docx"


def _docx_payload(*, document: bytes = b"<w:document/>") -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", b"<Types/>")
        archive.writestr("_rels/.rels", b"<Relationships/>")
        archive.writestr("word/document.xml", document)
    return output.getvalue()


def _real_docx_payload(text: str) -> bytes:
    output = io.BytesIO()
    document = Document()
    document.add_heading("Common Crawl DOCX", level=1)
    document.add_paragraph(text)
    document.save(output)
    return output.getvalue()


def _real_docx_with_table() -> bytes:
    output = io.BytesIO()
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Quarter"
    table.cell(0, 1).text = "Revenue"
    table.cell(1, 0).text = "Q1"
    table.cell(1, 1).text = "$42"
    document.save(output)
    return output.getvalue()


def _index_row(payload: bytes) -> dict[str, object]:
    return {
        "url": URL,
        "fetch_status": 200,
        "content_mime_type": DOCX_MIME_TYPE,
        "content_mime_detected": "application/zip",
        "content_digest": content_digest(payload),
        "content_truncated": None,
        "warc_filename": "crawl-data/test.warc.gz",
        "warc_record_offset": 42,
        "warc_record_length": 100,
        "warc_record_id": RECORD_ID,
    }


def _fetched(payload: bytes) -> FetchedCommonCrawlRecord:
    row = _index_row(payload)
    selection = DocxRecordSelector().select(row)
    assert selection is not None
    return FetchedCommonCrawlRecord(
        indexed_record=main_record_from_index_row(row, crawl_id=CRAWL_ID),
        selection=selection,
        observed_record=CommonCrawlWarcRecord(
            payload=payload,
            payload_digest=content_digest(payload),
            warc_record_id=RECORD_ID,
            target_url=URL,
            http_status=200,
            http_content_type=DOCX_MIME_TYPE,
            warc_date="2026-07-21T21:48:44Z",
            identified_payload_type="application/zip",
        ),
    )


def _config(**kwargs: object) -> CommonCrawlDocxConfig:
    return CommonCrawlDocxConfig(
        name="test-docx",
        sources=(
            CommonCrawlSource(
                crawl_id=CRAWL_ID,
                index_kind=CommonCrawlIndexKind.MAIN,
                paths_manifest_url="https://example.com/index.paths.gz",
            ),
        ),
        **kwargs,
    )


@dataclass(frozen=True)
class _Extractor:
    text: str
    version: str = "test-extractor-v1"

    def extract(self, payload: bytes) -> ExtractedDocument:
        assert payload.startswith(b"PK")
        return ExtractedDocument(
            text=self.text,
            word_count=len(self.text.split()),
            table_count=2,
            image_count=1,
            language_blocks=(DocumentBlock(self.text, is_table=False),),
        )


@dataclass(frozen=True)
class _LanguageDetector:
    scores: dict[str, float]
    version: str = "test-detector-v1"

    def score(self, text: str) -> dict[str, float]:
        assert text
        return self.scores


@dataclass(frozen=True)
class _MissingBlocksExtractor:
    version: str = "missing-blocks-v1"

    def extract(self, payload: bytes) -> ExtractedDocument:
        return ExtractedDocument(text="Extracted text", word_count=2, table_count=0, image_count=0, language_blocks=())


@pytest.mark.parametrize(
    ("row", "reason"),
    [
        ({"fetch_status": 200, "content_mime_type": DOCX_MIME_TYPE, "url": "https://example.com/a"}, "declared_mime"),
        ({"fetch_status": 200, "content_mime_type": "application/octet-stream", "url": URL}, "url_suffix"),
        (
            {"fetch_status": 200, "content_mime_detected": DOCX_MIME_TYPE, "url": "https://example.com/a"},
            "detected_mime",
        ),
    ],
)
def test_selector_preserves_selection_reason(row: dict[str, object], reason: str) -> None:
    selection = DocxRecordSelector().select(row)

    assert selection is not None
    assert selection.metadata["selection_reason"] == reason


def test_selector_rejects_failed_and_truncated_rows() -> None:
    selector = DocxRecordSelector()

    assert selector.select({"fetch_status": 404, "content_mime_type": DOCX_MIME_TYPE, "url": URL}) is None
    assert (
        selector.select(
            {"fetch_status": 200, "content_truncated": "length", "content_mime_type": DOCX_MIME_TYPE, "url": URL}
        )
        is None
    )


def test_validate_docx_accepts_required_office_members() -> None:
    validate_docx(_docx_payload(), maximum_entries=10, maximum_uncompressed_bytes=1024)


@pytest.mark.parametrize("payload", [b"not a zip", _docx_payload(document=b"x")[:-8]])
def test_validate_docx_rejects_malformed_payload(payload: bytes) -> None:
    with pytest.raises(InvalidDocxError):
        validate_docx(payload, maximum_entries=10, maximum_uncompressed_bytes=1024)


def test_validate_docx_enforces_entry_and_size_limits() -> None:
    payload = _docx_payload(document=b"x" * 100)

    with pytest.raises(InvalidDocxError, match="entries"):
        validate_docx(payload, maximum_entries=1, maximum_uncompressed_bytes=1024)
    with pytest.raises(InvalidDocxError, match="expands"):
        validate_docx(payload, maximum_entries=10, maximum_uncompressed_bytes=10)


def test_docling_extractor_reads_real_docx_fixture() -> None:
    extracted = DoclingDocxExtractor().extract(_real_docx_payload("A small document extracted by Docling."))

    assert "Common Crawl DOCX" in extracted.text
    assert "A small document extracted by Docling." in extracted.text


def test_docling_extractor_preserves_table_text_without_markdown_padding() -> None:
    extracted = DoclingDocxExtractor().extract(_real_docx_with_table())

    assert "Quarter | Revenue" in extracted.text
    assert "Q1 | $42" in extracted.text
    assert "---" not in extracted.text
    assert extracted.table_count == 1
    assert extracted.language_blocks[-1].is_table
    assert "Quarter | Revenue" in extracted.language_blocks[-1].text


def test_lingua_detector_identifies_multilingual_fixture() -> None:
    french = "Ceci est un document français avec suffisamment de mots pour identifier correctement sa langue. " * 3

    scores = LinguaLanguageDetector().score(french)

    assert max(scores, key=scores.__getitem__) == "fr"
    assert scores["fr"] > 0.5


def test_extracted_record_preserves_discovery_and_warc_provenance() -> None:
    fetched = fetched_docx_record(_fetched(_docx_payload()))
    output = extracted_docx_record(
        fetched,
        extractor=_Extractor("Bonjour tout le monde."),
        maximum_zip_entries=10,
        maximum_uncompressed_bytes=1024,
    )

    assert output["source_id"] == RECORD_ID
    assert output["crawl_id"] == CRAWL_ID
    assert output["warc_record_offset"] == 42
    assert output["index_status"] == 200
    assert output["index_content_type"] == DOCX_MIME_TYPE
    assert output["selection_reason"] == DocxSelectionReason.DECLARED_MIME.value
    assert "language" not in output
    assert output["table_count"] == 2
    assert output["language_blocks"] == [{"start": 0, "stop": 22, "is_table": False}]


def test_document_local_extraction_failure_is_skipped() -> None:
    output = process_fetched_docx(
        fetched_docx_record(_fetched(_docx_payload())),
        extractor=_Extractor("  \n"),
        maximum_zip_entries=10,
        maximum_uncompressed_bytes=1024,
    )

    assert output is None


def test_extracted_record_rejects_missing_language_blocks() -> None:
    with pytest.raises(DocxExtractionError, match="no language blocks"):
        extracted_docx_record(
            fetched_docx_record(_fetched(_docx_payload())),
            extractor=_MissingBlocksExtractor(),
            maximum_zip_entries=10,
            maximum_uncompressed_bytes=1024,
        )


def test_pipeline_separates_fetch_extraction_language_and_normalization() -> None:
    discovery, plan, fetch, extraction, language, normalized = common_crawl_docx_steps(
        _config(index_batch_rows=7, max_workers=11),
        extractor=_Extractor("text", version="extractor-v7"),
        language_detector=_LanguageDetector({"en": 1.0}, version="detector-v3"),
    )

    assert discovery in plan.deps
    assert plan in fetch.deps
    assert fetch in extraction.deps
    assert extraction in language.deps
    assert language in normalized.deps
    assert "index_batch_rows" not in discovery.hash_attrs
    assert "max_workers" not in discovery.hash_attrs
    assert extraction.hash_attrs["extractor"] == "extractor-v7"
    assert "language_detector" not in extraction.hash_attrs
    assert language.hash_attrs["language_detector"] == "detector-v3"


def test_language_chunks_clean_markdown_gate_noise_and_cap_table_weight() -> None:
    chunks = language_chunks(
        (
            DocumentBlock(
                "# Heading with enough alphabetic content to pass the configured gate.\n\n1234 !!!",
                is_table=False,
            ),
            DocumentBlock(
                "French table content with many alphabetic values repeated repeatedly repeatedly",
                is_table=True,
            ),
        ),
        minimum_alpha_bytes=10,
    )

    assert len(chunks) == 2
    assert "#" not in chunks[0].text
    assert "|" not in chunks[1].text
    assert chunks[1].is_table
    assert chunks[1].weight == chunks[1].alpha_bytes


def test_language_chunks_gate_low_alphabetic_ratio() -> None:
    chunks = language_chunks(
        (DocumentBlock("alphabeticcontent" * 4 + " 1234567890" * 30, is_table=False),),
        minimum_alpha_bytes=10,
        minimum_alpha_ratio=0.5,
    )

    assert chunks == []


def test_language_chunks_cap_total_table_weight_across_all_chunks() -> None:
    prose = "prosecontent " * 6
    blocks = (
        DocumentBlock(prose, is_table=False),
        DocumentBlock("tablecontent " * 200, is_table=True),
    )

    detection = detect_document_language(
        blocks,
        detector=_TableLanguageDetector(),
        chunk_chars=100,
        sample_chunks=12,
        minimum_alpha_bytes=10,
        maximum_table_alpha_bytes=50,
    )

    prose_alpha_bytes = sum(character.isalpha() for character in prose)
    assert detection.chunks_total > detection.chunks_scored
    assert detection.distribution["fr"] == pytest.approx(50 / (50 + prose_alpha_bytes))


@dataclass(frozen=True)
class _TableLanguageDetector:
    version: str = "table-detector-v1"

    def score(self, text: str) -> dict[str, float]:
        return {"fr": 1.0, "en": 0.0} if "tablecontent" in text else {"fr": 0.0, "en": 1.0}


@dataclass(frozen=True)
class _TextLanguageDetector:
    version: str = "text-detector-v1"

    def score(self, text: str) -> dict[str, float]:
        return {"fr": 1.0, "en": 0.0} if "français" in text else {"fr": 0.0, "en": 1.0}


def test_document_language_uses_alpha_byte_weighted_chunk_scores() -> None:
    english = "English words form a substantially longer body section for weighting. " * 8
    french = "Texte français suffisamment long pour passer le filtre avec encore plusieurs mots alphabétiques."

    detection = detect_document_language(
        (DocumentBlock(f"{english}\n\n{french}", is_table=False),), detector=_TextLanguageDetector()
    )

    assert detection.language == "en"
    assert detection.distribution["en"] > detection.distribution["fr"]
    assert detection.chunks_scored == 2


def test_document_language_limits_persisted_distribution() -> None:
    detector = _LanguageDetector({f"language-{index}": 1.0 / (index + 1) for index in range(10)})

    detection = detect_document_language(
        (DocumentBlock("English alphabetic content long enough to pass the language gate.", is_table=False),),
        detector=detector,
        distribution_top_k=3,
    )

    assert list(detection.distribution) == ["language-0", "language-1", "language-2"]


def test_document_language_marks_low_confidence_prediction_unknown() -> None:
    detection = detect_document_language(
        (DocumentBlock("English alphabetic content long enough to pass the language gate.", is_table=False),),
        detector=_LanguageDetector({"en": 0.4, "fr": 0.35}),
        minimum_score=0.5,
    )

    assert detection.language == "unknown"
    assert detection.score == 0.4


@dataclass
class _RecordingLanguageDetector:
    seen: list[str]
    version: str = "recording-detector-v1"

    def score(self, text: str) -> dict[str, float]:
        self.seen.append(text)
        return {"en": 1.0}


def test_document_language_long_input_samples_across_entire_document() -> None:
    blocks = tuple(DocumentBlock(f"section{index} " + "alphabetic content " * 8, is_table=False) for index in range(20))
    detector = _RecordingLanguageDetector([])

    detection = detect_document_language(blocks, detector=detector, sample_chunks=12)

    assert detection.chunks_total == 20
    assert detection.chunks_scored == 12
    assert "section0" in detector.seen[0]
    assert "section19" in detector.seen[-1]


def test_language_identification_removes_structural_blocks_from_output() -> None:
    extracted = extracted_docx_record(
        fetched_docx_record(_fetched(_docx_payload())),
        extractor=_Extractor("English body content with enough alphabetic letters to pass language identification."),
        maximum_zip_entries=10,
        maximum_uncompressed_bytes=1024,
    )

    output = language_identified_docx_record(extracted, detector=_TextLanguageDetector(), config=_config())

    assert output["language"] == "en"
    assert "language_blocks" not in output


def test_sample_report_reads_selection_reason_from_shared_discovery_metadata() -> None:
    source = _config().sources[0]
    candidates = [
        {"selection_metadata": json.dumps({"selection_reason": DocxSelectionReason.DECLARED_MIME.value})},
        {"selection_metadata": json.dumps({"selection_reason": DocxSelectionReason.DECLARED_MIME.value})},
        {"selection_metadata": json.dumps({"selection_reason": DocxSelectionReason.URL_SUFFIX.value})},
    ]
    extracted = [
        {
            "selection_reason": DocxSelectionReason.DECLARED_MIME.value,
            "url": URL,
            "language": "en",
            "word_count": 4,
            "table_count": 1,
            "text": "A short extracted document.",
        }
    ]

    markdown, examples = sample_report_markdown(
        source=source,
        candidate_rows=candidates,
        extracted_rows=extracted,
        normalized_rows=extracted,
        stage_counters={"extraction/common_crawl_docx/invalid_files": 2},
        examples_per_reason=1,
    )

    assert "| `declared_mime` | 2 | 1 | 50.0% |" in markdown
    assert "| `url_suffix` | 1 | 0 | 0.0% |" in markdown
    assert "`extraction/common_crawl_docx/invalid_files`: 2" in markdown
    assert examples == [
        {
            "selection_reason": DocxSelectionReason.DECLARED_MIME.value,
            "url": URL,
            "language": "en",
            "word_count": 4,
            "table_count": 1,
            "excerpt": "A short extracted document.",
        }
    ]
